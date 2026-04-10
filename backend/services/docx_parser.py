"""DOCX Parser — extracts structured content from אור הפרשה DOCX files.

Iterates doc.element.body (not doc.paragraphs) to preserve correct ordering
of paragraphs, tables, and images. Handles textbox dedup, nikud stripping,
bold detection, and word-search table extraction.
"""

import re
import zipfile
import xml.etree.ElementTree as ET
from docx import Document


def strip_nikud(text: str) -> str:
    """Remove Hebrew vowel marks (nikud) for matching."""
    return re.sub(r'[\u0591-\u05C7]', '', text)


def is_entirely_bold(para) -> bool:
    """Check if all non-empty runs in a paragraph are bold."""
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def render_runs(para) -> str:
    """Render paragraph runs preserving bold formatting."""
    parts = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        from html import escape
        t = escape(t)
        parts.append(f'<strong>{t}</strong>' if run.bold else t)
    return ''.join(parts)


def para_has_image(para) -> str | None:
    """Check if paragraph contains an image, return rId or None."""
    ns = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
    for elem in para._element.iter():
        if elem.tag.endswith('}blip'):
            return elem.get(f'{ns}embed')
    return None


def extract_textbox_texts(para) -> list[str]:
    """Extract text from textboxes (floating shapes) in a paragraph."""
    texts = []
    ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for txbx in para._element.iter():
        if txbx.tag.endswith('}txbxContent'):
            parts = []
            for p in txbx.iter(f'{ns_w}t'):
                if p.text:
                    parts.append(p.text)
            text = ''.join(parts).strip()
            if text:
                texts.append(text)
    return texts


def extract_table_data(table) -> list[list[str]]:
    """Extract table data as list of rows, each row a list of cell texts."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cells.append(cell.text.strip())
        rows.append(cells)
    return rows


def detect_section_break(text: str) -> bool:
    """Detect if text marks the teacher section start."""
    stripped = strip_nikud(text).strip()
    teacher_markers = ['למלמדים', 'למורים', 'דגשים למלמד', 'הדגשים', 'דגשי הוראה']
    for marker in teacher_markers:
        if marker in stripped:
            return True
    return False


def parse_docx(docx_path: str) -> dict:
    """Parse a DOCX file and return structured content.

    Returns dict with:
        parsha_name: str — detected parsha name
        subtitle: str — detected subtitle
        dedication: str — dedication text if found
        children_content: list — content blocks for children section
        teacher_content: list — content blocks for teacher section
        images: dict — {rId: local_path} for images
    """
    doc = Document(docx_path)

    # Extract images map (rId → target path in ZIP)
    image_rels = {}
    with zipfile.ZipFile(docx_path) as z:
        try:
            rels_xml = z.read("word/_rels/document.xml.rels")
            root = ET.fromstring(rels_xml)
            for r in root:
                if "image" in r.get("Type", "").lower():
                    image_rels[r.get("Id")] = r.get("Target")
        except KeyError:
            pass

    parsha_name = ""
    subtitle = ""
    dedication = ""
    children_content = []
    teacher_content = []
    in_teacher_section = False
    seen_textbox_texts = set()

    para_count = 0
    for elem in doc.element.body:
        tag = elem.tag.split('}')[-1]

        if tag == 'p':
            para = doc.paragraphs[para_count]
            para_count += 1
            text = para.text.strip()
            text_stripped = strip_nikud(text)

            if not text and not para_has_image(para):
                # Check for textbox content
                tb_texts = extract_textbox_texts(para)
                for tb in tb_texts:
                    if tb not in seen_textbox_texts:
                        seen_textbox_texts.add(tb)
                        target = teacher_content if in_teacher_section else children_content
                        target.append({"type": "textbox", "text": tb})
                continue

            # Detect parsha name (usually first meaningful line with אור הפרשה)
            if not parsha_name and 'פרשת' in text_stripped:
                parsha_name = text
                continue

            # Detect subtitle (second line, often the theme)
            if parsha_name and not subtitle and text and len(text) < 100:
                if para_count <= 5:
                    subtitle = text
                    continue

            # Detect dedication (contains לע"נ or הוקדש)
            if ('לע"נ' in text or 'הוקדש' in text_stripped or 'לעילוי' in text_stripped) and not dedication:
                dedication = text
                continue

            # Detect section break to teacher content
            if detect_section_break(text):
                in_teacher_section = True

            # Check for image
            img_rid = para_has_image(para)

            # Build content block
            target = teacher_content if in_teacher_section else children_content
            is_heading = is_entirely_bold(para) and len(text) < 90

            if img_rid:
                target.append({"type": "image", "rId": img_rid})

            if text:
                if is_heading:
                    target.append({"type": "heading", "html": render_runs(para), "text": text})
                else:
                    target.append({"type": "paragraph", "html": render_runs(para), "text": text})

        elif tag == 'tbl':
            # Find matching table object
            for table in doc.tables:
                if table._tbl is elem:
                    table_data = extract_table_data(table)
                    target = teacher_content if in_teacher_section else children_content
                    target.append({"type": "table", "data": table_data})
                    break

    return {
        "parsha_name": strip_nikud(parsha_name),
        "subtitle": subtitle,
        "dedication": dedication,
        "children_content": children_content,
        "teacher_content": teacher_content,
        "image_rels": image_rels,
        "docx_path": docx_path,
    }
